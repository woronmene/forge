from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from pathlib import Path

ROOT = Path('/Users/woron/Documents/dev/forge')
OUT = ROOT / 'docs' / 'forge-backend-integration-status-report.docx'
SHOT = ROOT / 'docs' / 'artifacts' / 'screenshots'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color='DADCE0'):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        element = tc_borders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tc_borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '6')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def format_cell(cell, fill=None, bold=False, font_size=10, color='000000'):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(font_size)
            r.font.bold = bold
            r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)
    if fill:
        set_cell_shading(cell, fill)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_after = Pt(4)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.name = 'Arial'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(85, 85, 85)
        p2.paragraph_format.space_after = Pt(12)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.bold = False
    r.font.color.rgb = RGBColor(0, 0, 0) if level == 1 else RGBColor(67, 67, 67)
    r.font.size = Pt(16 if level == 1 else 13)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)


def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, width in enumerate(widths):
        table.columns[i].width = width
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        format_cell(hdr[i], fill='F1F3F4', bold=True, font_size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            format_cell(cells[i], font_size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_image(doc, path, caption):
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.0))
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(85, 85, 85)
        cp.paragraph_format.space_after = Pt(8)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

styles = doc.styles
for style_name in ['Normal', 'List Bullet']:
    style = styles[style_name]
    style.font.name = 'Arial'
    style.font.size = Pt(11)

add_title(doc, 'Forge Backend Integration Status Report', 'Simple testing guide for the current frontend-to-backend implementation')
add_para(doc, 'This report shows what is already connected to real backend endpoints, what is only partially connected, and what still needs to be provided by the backend team before every component can become fully functional.')

add_heading(doc, '1. What you can test right now')
rows = [
    ['Admin sign-in', 'Live', 'Real admin auth works with the provided development credential set.'],
    ['Content library', 'Live', 'Loads real media assets from staging.'],
    ['Analytics', 'Live', 'Loads real overview metrics from staging.'],
    ['Engagement overview', 'Live', 'Loads real engagement summary payload from staging.'],
    ['Wallet overview', 'Live / partial', 'Overview loads; transaction history is not available on this staging contract.'],
    ['Processing queue', 'Live / partial', 'Uses real media asset list and retry actions, but not a true worker queue endpoint.'],
    ['Series page', 'Partial', 'Shows backend-backed catalog inventory, but not full series editor data.'],
    ['Albums & mixes page', 'Partial', 'Shows backend-backed catalog inventory, but not full editor data.'],
    ['Upload page', 'Partial', 'UI exists, but final authenticated upload flow is not fully implemented.'],
    ['Users page', 'Partial', 'User detail flows exist, but full directory listing is still blocked.'],
]
add_table(doc, ['Area', 'Status', 'What it means'], rows, [Inches(1.6), Inches(1.1), Inches(3.8)])

add_heading(doc, '2. Page-by-page testing notes')
add_bullets(doc, [
    'Sign in: should take you straight into the dashboard.',
    'Content library: should show 2 staging assets right now and mark them as Queued.',
    'Analytics: should load real totals from the analytics overview endpoint, even if many values are currently 0.',
    'Engagement overview: should load daily active users, monthly active users, and registrations from the live engagement endpoint.',
    'Wallet activity: should load overview cards and show a clear notice that transaction history is not available on this staging environment.',
    'Processing queue: should show the same real media assets as backend-backed processing rows.',
])

add_image(doc, SHOT / 'content-library.png', 'Content library: real media asset list is loading from staging.')
add_image(doc, SHOT / 'analytics.png', 'Analytics: overview cards are connected to the live analytics service.')

doc.add_page_break()
add_image(doc, SHOT / 'engagement.png', 'Engagement overview: connected to the live user engagement endpoint.')
add_image(doc, SHOT / 'wallet.png', 'Wallet activity: overview is live, transaction history is still unavailable on this environment.')
add_image(doc, SHOT / 'processing.png', 'Processing queue: driven from the real media asset list and retry-processing capability.')

add_heading(doc, '3. What is still partial or blocked')
rows2 = [
    ['Upload flow', 'Need exact final upload + complete-upload payload sequence, plus any create-entity mapping for movie/series/album/mix/trailer.'],
    ['Users directory', 'Need a confirmed list-users endpoint with filters and pagination.'],
    ['Series editor', 'Need a list/detail contract that returns series with nested seasons and episodes ready for drawer editing.'],
    ['Albums & mixes editor', 'Need a list/detail contract for albums and mixes, plus cue-point payload confirmation for mixes.'],
    ['Wallet transactions', 'Need the actual transaction-history route and request/response shape for this staging environment.'],
    ['Wallet provider health / card issuance', 'Need dedicated endpoints or a documented payload source; current UI only has placeholders for these.'],
    ['Processing telemetry', 'Need a real job queue / worker status endpoint if the design is meant to show pipeline-level processing, not just asset status.'],
    ['Content summaries', 'The internal media summary and catalog summary routes previously assumed in the UI are not present on this staging environment.'],
]
add_table(doc, ['Component / module', 'What backend still needs to be provided'], rows2, [Inches(2.0), Inches(4.5)])

add_heading(doc, '4. Endpoint mapping that is already confirmed')
add_bullets(doc, [
    'Admin login: POST /internal/v1/admin/auth/login',
    'Admin refresh: POST /internal/v1/admin/auth/refresh',
    'Media asset list: GET /v1/media/assets',
    'Analytics overview: GET /v1/analytics/media/overview',
    'Engagement overview: GET /v1/analytics/users/engagement',
    'Wallet overview: GET /v1/analytics/wallet/overview',
    'Retry processing: media retry routes are wired from the processing screen.',
])

add_heading(doc, '5. What to report back to the backend team')
add_bullets(doc, [
    'The admin sign-in flow works and has been tested successfully with a superadmin credential.',
    'The content library, analytics, engagement overview, wallet overview, and processing queue are already connected to real staging endpoints.',
    'The wallet transaction history endpoint used in the earlier collection does not resolve on the current staging environment and needs the correct route or updated contract.',
    'The users page still needs a real directory list endpoint.',
    'The upload flow still needs the exact authenticated upload contract for final implementation.',
    'The series and albums screens still need richer entity list/detail payloads to support the full drawer components in the design.',
])

add_heading(doc, '6. Practical testing order')
add_bullets(doc, [
    '1. Test sign-in.',
    '2. Open content library and confirm real rows appear.',
    '3. Open analytics and engagement and confirm the pages load without auth errors.',
    '4. Open wallet activity and confirm the overview loads and the transaction-history gap notice appears.',
    '5. Open processing queue and confirm the same staging assets appear there.',
    '6. Report the blocked backend items listed above for the remaining dummy or partially implemented components.',
])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
